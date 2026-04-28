"""
Genera dos archivos para el protocolo de visitas PYGLARA:
  1. PYGLARA_Protocolo_Visitas.docx  — Guia de campo para empleados (para imprimir)
  2. PYGLARA_Registro_Visitas.xlsx   — Formulario de seguimiento de visitas
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────────────────────
# COLORS
# ─────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1A, 0x2E, 0x4A)   # dark navy blue
GOLD     = RGBColor(0xC8, 0x9B, 0x2F)   # warm gold
LGRAY    = RGBColor(0xF4, 0xF4, 0xF4)   # light gray background
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
RED_WARN = RGBColor(0xC0, 0x20, 0x20)

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Set table cell background via XML shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1, color=None, center=False):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if color:
        run.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.25)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_step_box(doc, step_num, title, bullets):
    """Add a shaded step block with numbered heading and bullet points."""
    # Step heading paragraph (shaded)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"  PASO {step_num} — {title}  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    # background via paragraph XML (approximation via highlight is limited; use a table row)

    # Use a 1-cell table for colored header
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.text = f"PASO {step_num} — {title}"
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].font.color.rgb = WHITE
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_cell_bg(cell, '1A2E4A')

    # Remove the loose paragraph we added
    p._element.getparent().remove(p._element)

    for bullet in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.left_indent  = Inches(0.3)
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after  = Pt(1)
        if isinstance(bullet, tuple):
            # (bold_part, rest)
            r1 = bp.add_run(bullet[0])
            r1.bold = True
            r1.font.size = Pt(10)
            if len(bullet) > 1:
                r2 = bp.add_run(bullet[1])
                r2.font.size = Pt(10)
        else:
            run = bp.add_run(bullet)
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacer


# ─────────────────────────────────────────────────────────────
# WORD DOCUMENT
# ─────────────────────────────────────────────────────────────
def build_word():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── TITLE BLOCK ──────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PYGLARA")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("Prensados y Galvanizados de Lara, S.A.")
    r2.font.size = Pt(11)
    r2.font.color.rgb = GOLD
    r2.bold = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("PROTOCOLO DE VISITAS A CLIENTES — ZONA INDUSTRIAL I")
    r3.bold = True
    r3.font.size = Pt(13)
    r3.font.color.rgb = NAVY

    t4 = doc.add_paragraph()
    t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = t4.add_run("Guia de campo para empleados | Abril 2026")
    r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # ── OBJETIVO ─────────────────────────────────────────────
    tbl_obj = doc.add_table(rows=1, cols=1)
    tbl_obj.style = 'Table Grid'
    c = tbl_obj.rows[0].cells[0]
    c.paragraphs[0].clear()
    set_cell_bg(c, 'F4F4F4')
    p_obj = c.paragraphs[0]
    p_obj.paragraph_format.space_before = Pt(4)
    p_obj.paragraph_format.space_after  = Pt(4)
    r_lbl = p_obj.add_run("OBJETIVO:  ")
    r_lbl.bold = True
    r_lbl.font.color.rgb = NAVY
    r_lbl.font.size = Pt(10)
    r_txt = p_obj.add_run(
        "Retomar el contacto con clientes historicos y empresas vecinas para informarles "
        "de la proxima reactivacion de PYGLARA y levantar intenciones de pedidos antes de "
        "que llegue el zinc."
    )
    r_txt.font.size = Pt(10)

    doc.add_paragraph()

    # ── CLIENTES PRIORITARIOS TABLE ───────────────────────────
    p_sec = doc.add_paragraph()
    rs = p_sec.add_run("CLIENTES PRIORITARIOS — ZONA INDUSTRIAL I (visita en persona)")
    rs.bold = True
    rs.font.size = Pt(11)
    rs.font.color.rgb = NAVY

    headers = ["#", "Empresa", "Que producen / necesitan", "Telefono"]
    rows_data = [
        ["1", "SASGO",
         "Torres y postes electricos para CORPOELEC. ALTA necesidad de galvanizado.",
         "+58 412-536-3346"],
        ["2", "Industrias Marullo, S.A.",
         "Maquinaria agroindustrial y estructuras de acero. Fundada 1955. Muy activa.",
         "+58 424-514-3859"],
        ["3", "GEDISA",
         "Distribuidor electrico de 60 anos. Reabrieron enero 2026.",
         "+58 414-503-2546"],
        ["4", "CORPOELEC Lara — Alm. N01",
         "Almacen regional del Estado. Compraron 20.000 varillas de tierra a PYGLARA (2004-2005).",
         "0251-239-4050"],
        ["5", "MANPEG",
         "Fabrica postes hexagonales y electricos hasta 40m.",
         "+58 424-561-2759"],
        ["6", "GEMACA",
         "Tuberia PEAD. Necesidad baja — verificar.",
         "+58 251-237-6922"],
    ]

    tbl = doc.add_table(rows=1 + len(rows_data), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    widths = [Cm(1), Cm(4.5), Cm(9), Cm(3.5)]
    for i, (hdr, w) in enumerate(zip(headers, widths)):
        cell = tbl.rows[0].cells[i]
        cell.width = w
        set_cell_bg(cell, '1A2E4A')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)

    # Data rows
    for ri, row_vals in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        bg = 'F4F4F4' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
                run.font.color.rgb = NAVY

    doc.add_paragraph()

    # ── STEPS ────────────────────────────────────────────────
    add_step_box(doc, 1, "PREPARACION ANTES DE SALIR (10 min en planta)", [
        ("Que llevar:", ""),
        "[ ]  Tarjeta de presentacion de PYGLARA (o hoja membretada con nombre, cargo, telefono)",
        "[ ]  Formulario de Registro de Visita impreso — UNA copia por empresa",
        "[ ]  Boligrafo",
        "[ ]  Telefono cargado",
        ("Repasar:", " nombre de la empresa, que producen, y el mensaje clave (Paso 3)",),
    ])

    add_step_box(doc, 2, "LLEGADA A LA EMPRESA", [
        'Presentate en recepcion o con el vigilante:',
        '"Buenos dias. Mi nombre es [TU NOMBRE], soy de Prensados y Galvanizados de Lara — '
        'PYGLARA, estamos ubicados aqui mismo en la Zona Industrial. '
        'Vengo a hablar con el encargado de compras o el gerente de planta."',
        'Si no hay nadie: solicita nombre del responsable, cuando regresar y deja tarjeta.',
        'Si te reciben: pide 5 minutos para explicar el motivo de la visita.',
    ])

    add_step_box(doc, 3, "EL MENSAJE CLAVE (lo que debes decir)", [
        ("VERSION CORTA (2 min):  ", ""),
        '"PYGLARA — el galvanizador de la Zona Industrial — esta pronto a reactivarse. '
        'Paramos en 2015 por falta de zinc, pero ya estamos gestionando la compra. '
        'Queremos contactar a clientes para coordinar pedidos. '
        '¿Tienen alguna necesidad de galvanizado o varillas de puesta a tierra cobreadas?"',
        ("Si preguntan por precios:  ",
         '"Trabajamos por kilo galvanizado, rango de mercado $1.50–$1.70/kg. '
         'Para precio formal necesitamos conocer tipo de pieza y volumen."'),
        ("Si preguntan cuando arrancan:  ",
         '"El zinc tarda ~60 dias en llegar, luego 3 semanas de puesta en marcha. '
         'Estamos en gestion ahora mismo."'),
        ("Si ya tienen otro proveedor:  ",
         '"Lo entendemos. Somos uno de solo 2 galvanizadores activos en Venezuela. '
         'Si necesitan capacidad adicional, nos llaman."'),
    ])

    add_step_box(doc, 4, "OBJETIVO: CONSEGUIR UNA REUNION", [
        ("NO es cerrar un pedido hoy.  ", "El objetivo es:"),
        "a)  Saber si tienen necesidad (si / no / cuanto)",
        "b)  Conseguir nombre y datos del responsable de compras o mantenimiento",
        "c)  Agendar reunion de seguimiento — en persona, por telefono o WhatsApp",
        ('"¿Podriamos coordinar 20 minutos con usted o la persona de compras '
         'esta semana?"  ', ""),
        '"¿Le parece si le escribo por WhatsApp? ¿Cual es el mejor numero?"',
    ])

    add_step_box(doc, 5, "ANTES DE RETIRARSE", [
        "Agradece el tiempo dedicado.",
        "Deja tarjeta o nota con: nombre PYGLARA, tu nombre, telefono de contacto.",
        "Escribe en la nota:  \"PYGLARA — proxima reactivacion 2026\"",
        "Llena el Formulario de Registro antes de ir a la siguiente empresa.",
    ])

    add_step_box(doc, 6, "LLENAR EL FORMULARIO DE REGISTRO", [
        "Llena UNA copia por cada empresa visitada.",
        "Completa todos los campos — aunque la respuesta sea 'no habia nadie'.",
        "Entrega todos los formularios a gerencia al final del dia.",
    ])

    add_step_box(doc, 7, "MENSAJE DE SEGUIMIENTO (dentro de 24 horas)", [
        ("WhatsApp:  ", ""),
        '"Buenas [tardes], [NOMBRE]. Le escribo de PYGLARA — estuvimos por su empresa hoy. '
        'Queremos coordinar una reunion breve para presentarles nuestras capacidades de '
        'galvanizado en caliente y varillas de tierra cobreadas. '
        'Estamos en proceso de reactivacion este 2026. '
        '¿Tiene disponibilidad esta semana o la proxima? Gracias."',
        ("Correo — Asunto:  ",
         '"PYGLARA — Galvanizado en Caliente | Reactivacion 2026 | Solicitud de Reunion"'),
        "El texto completo del correo lo tiene el coordinador comercial — solicitarlo si se necesita.",
    ])

    add_step_box(doc, 8, "ENTREGA A GERENCIA AL FINAL DEL DIA", [
        "Reune todos los formularios llenados.",
        "Entregalos fisicamente al coordinador o gerencia.",
        "Comenta verbalmente los 2-3 puntos mas importantes de lo que escuchaste.",
    ])

    # ── RESUMEN TABLE ─────────────────────────────────────────
    p_res = doc.add_paragraph()
    rr = p_res.add_run("RESUMEN RAPIDO")
    rr.bold = True
    rr.font.size = Pt(11)
    rr.font.color.rgb = NAVY

    summary = [
        ("Antes de salir",      "Tarjeta/hoja + formulario impreso + boligrafo + telefono cargado"),
        ("Al llegar",           "Pide hablar con encargado de compras o gerente de planta"),
        ("Mensaje clave",       "\"PYGLARA se reactiva en 2026 — queremos coordinar pedidos\""),
        ("Objetivo",            "Conseguir nombre de contacto y agendar reunion"),
        ("Si no hay nadie",     "Deja tarjeta, anota nombre y cuando regresar"),
        ("Despues de visitar",  "Llena formulario, envia WhatsApp en menos de 24 horas"),
        ("Final del dia",       "Entrega todos los formularios a gerencia"),
    ]

    ts = doc.add_table(rows=len(summary), cols=2)
    ts.style = 'Table Grid'
    ts.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (label, desc) in enumerate(summary):
        bg = '1A2E4A' if ri == 0 else ('F0F4FA' if ri % 2 == 0 else 'FFFFFF')
        c0 = ts.rows[ri].cells[0]
        c1 = ts.rows[ri].cells[1]
        c0.width = Cm(4)
        c1.width = Cm(13)
        set_cell_bg(c0, '1A2E4A' if ri % 2 == 0 else '223355')
        set_cell_bg(c1, 'F4F4F4' if ri % 2 == 0 else 'FFFFFF')
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.color.rgb = WHITE
        r0.font.size = Pt(9)
        r1 = c1.paragraphs[0].add_run(desc)
        r1.font.size = Pt(9)

    doc.add_paragraph()

    # ── FOOTER NOTE ───────────────────────────────────────────
    fn = doc.add_paragraph()
    fn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rfn = fn.add_run(
        "Documento interno PYGLARA — Abril 2026  |  "
        "Zona Industrial I, Barquisimeto, Estado Lara  |  Tel: +58 424 571 5349"
    )
    rfn.font.size = Pt(8)
    rfn.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    rfn.italic = True

    out_path = os.path.join(OUTPUT_DIR, "PYGLARA_Protocolo_Visitas.docx")
    doc.save(out_path)
    print(f"WORD guardado: {out_path}")


# ─────────────────────────────────────────────────────────────
# EXCEL WORKBOOK
# ─────────────────────────────────────────────────────────────
def xl_fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def xl_font(bold=False, size=10, color="000000", italic=False):
    return Font(bold=bold, size=size, color=color, italic=italic, name="Calibri")

def xl_border(style="thin"):
    side = Side(style=style, color="AAAAAA")
    return Border(left=side, right=side, top=side, bottom=side)

def xl_align(horizontal="left", vertical="center", wrap=False):
    return Alignment(horizontal=horizontal, vertical=vertical, wrap_text=wrap)


def build_excel():
    wb = openpyxl.Workbook()

    # ── SHEET 1: REGISTRO DE VISITAS ──────────────────────────
    ws = wb.active
    ws.title = "Registro de Visitas"
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A5"

    # -- Title rows
    ws.merge_cells("A1:O1")
    ws["A1"] = "PYGLARA — Prensados y Galvanizados de Lara, S.A."
    ws["A1"].font      = xl_font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill      = xl_fill("1A2E4A")
    ws["A1"].alignment = xl_align("center")
    ws.row_dimensions[1].height = 22

    ws.merge_cells("A2:O2")
    ws["A2"] = "REGISTRO DE VISITAS A CLIENTES — ZONA INDUSTRIAL I  |  Abril 2026"
    ws["A2"].font      = xl_font(bold=True, size=11, color="C89B2F")
    ws["A2"].fill      = xl_fill("1A2E4A")
    ws["A2"].alignment = xl_align("center")
    ws.row_dimensions[2].height = 18

    ws.merge_cells("A3:O3")
    ws["A3"] = (
        "Instrucciones: Llenar UNA fila por visita realizada. Entregar a gerencia al final del dia. "
        "Campos con * son obligatorios."
    )
    ws["A3"].font      = xl_font(size=9, italic=True, color="555555")
    ws["A3"].fill      = xl_fill("F4F4F4")
    ws["A3"].alignment = xl_align("center")
    ws.row_dimensions[3].height = 14

    ws.row_dimensions[4].height = 6  # spacer

    # -- Column headers (row 5)
    headers = [
        ("A", 12,  "Fecha*"),
        ("B", 6,   "#"),
        ("C", 22,  "Empresa Visitada*"),
        ("D", 20,  "Nombre Contacto"),
        ("E", 14,  "Cargo"),
        ("F", 16,  "Telefono / WA"),
        ("G", 22,  "Email"),
        ("H", 14,  "Habia persona disponible?"),
        ("I", 16,  "Interes Galvanizado*"),
        ("J", 16,  "Interes Varillas Tierra*"),
        ("K", 14,  "Vol. Estimado Galv. (TM/mes)"),
        ("L", 14,  "Vol. Estimado Varillas (uds/mes)"),
        ("M", 16,  "Seguimiento acordado"),
        ("N", 22,  "Fecha/Hora Seguimiento"),
        ("O", 30,  "Notas / Observaciones*"),
    ]

    for col_letter, width, label in headers:
        ws.column_dimensions[col_letter].width = width
        cell = ws[f"{col_letter}5"]
        cell.value     = label
        cell.font      = xl_font(bold=True, size=9, color="FFFFFF")
        cell.fill      = xl_fill("1A2E4A")
        cell.alignment = xl_align("center", wrap=True)
        cell.border    = xl_border()

    ws.row_dimensions[5].height = 32

    # -- Data rows (6 to 55 — 50 rows)
    options_si_no  = '"Si,No,No disponible"'
    options_interes = '"Si - necesidad activa,Tal vez,No,No se pudo determinar"'
    options_seguim  = '"Reunion agendada,Llamada acordada,WhatsApp enviado,Sin acuerdo,No aplica"'

    from openpyxl.worksheet.datavalidation import DataValidation

    dv_disponible = DataValidation(type="list", formula1=options_si_no,  showDropDown=False)
    dv_interes    = DataValidation(type="list", formula1=options_interes, showDropDown=False)
    dv_seguim     = DataValidation(type="list", formula1=options_seguim,  showDropDown=False)

    dv_disponible.sqref = "H6:H55"
    dv_interes.sqref    = "I6:J55"
    dv_seguim.sqref     = "M6:M55"

    ws.add_data_validation(dv_disponible)
    ws.add_data_validation(dv_interes)
    ws.add_data_validation(dv_seguim)

    for r in range(6, 56):
        ws.row_dimensions[r].height = 20
        bg = "F9FBFF" if r % 2 == 0 else "FFFFFF"
        for col_letter, _, _ in headers:
            cell = ws[f"{col_letter}{r}"]
            cell.fill      = xl_fill(bg)
            cell.font      = xl_font(size=9)
            cell.alignment = xl_align("left", "center", wrap=True)
            cell.border    = xl_border()

    # Pre-fill empresa column with priority clients for first 6 rows
    priority = [
        "SASGO",
        "Industrias Marullo, S.A.",
        "GEDISA",
        "CORPOELEC Lara — Almacen N01",
        "MANPEG",
        "GEMACA",
    ]
    for i, nombre in enumerate(priority):
        cell = ws[f"C{6+i}"]
        cell.value = nombre
        cell.font  = xl_font(size=9, color="1A2E4A")

    # ── SHEET 2: GUIA RAPIDA ──────────────────────────────────
    ws2 = wb.create_sheet("Guia Rapida")
    ws2.sheet_view.showGridLines = False

    ws2.merge_cells("A1:B1")
    ws2["A1"] = "PYGLARA — GUIA RAPIDA DE VISITAS"
    ws2["A1"].font      = xl_font(bold=True, size=13, color="FFFFFF")
    ws2["A1"].fill      = xl_fill("1A2E4A")
    ws2["A1"].alignment = xl_align("center")
    ws2.row_dimensions[1].height = 22
    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 55

    pasos = [
        ("",                  ""),
        ("PASO 1 — Preparacion",
         "Lleva: tarjeta PYGLARA, formulario impreso, boligrafo, telefono cargado"),
        ("PASO 2 — Llegada",
         "Pide hablar con encargado de compras o gerente de planta"),
        ("PASO 3 — Mensaje",
         '"PYGLARA se reactiva en 2026 — queremos coordinar pedidos de galvanizado y varillas de tierra"'),
        ("PASO 4 — Objetivo",
         "Conseguir nombre del contacto de compras y agendar reunion (NO cerrar pedido hoy)"),
        ("PASO 5 — Retirarse",
         "Deja tarjeta. Escribe: PYGLARA — proxima reactivacion 2026"),
        ("PASO 6 — Formulario",
         "Llena UNA fila en el Registro por cada empresa visitada"),
        ("PASO 7 — Seguimiento",
         "Envia WhatsApp o correo dentro de las 24 horas"),
        ("PASO 8 — Gerencia",
         "Entrega formularios completos al final del dia"),
        ("",                  ""),
        ("SI PREGUNTAN PRECIOS:",
         "$1.50 – $1.70 por kg galvanizado. Precio formal: requiere tipo de pieza y volumen"),
        ("SI PREGUNTAN CUANDO:",
         "Zinc tarda ~60 dias. Luego 3 semanas de puesta en marcha. En gestion ahora"),
        ("SI YA TIENEN PROVEEDOR:",
         "Solo hay 2 galvanizadores activos en Venezuela. Para capacidad adicional, nos llaman"),
        ("",                  ""),
        ("WHATSAPP DE SEGUIMIENTO:", ""),
        ("Texto:",
         '"Buenas [tardes], [NOMBRE]. Le escribo de PYGLARA — estuvimos hoy en su empresa. '
         'Queremos coordinar una reunion para presentarles nuestras capacidades de galvanizado '
         'y varillas de tierra cobreadas. Reactivacion 2026. '
         '¿Disponibilidad esta semana o la proxima? Gracias."'),
    ]

    for i, (label, desc) in enumerate(pasos, start=2):
        ws2.row_dimensions[i].height = 18 if desc else 8
        a = ws2.cell(row=i, column=1, value=label)
        b = ws2.cell(row=i, column=2, value=desc)
        is_step = label.startswith("PASO")
        is_qna  = label.endswith(":")
        a.font      = xl_font(bold=True,  size=9,  color="1A2E4A" if is_step else ("C89B2F" if is_qna else "333333"))
        a.fill      = xl_fill("EEF2F8" if is_step else "FFFFFF")
        b.font      = xl_font(size=9)
        b.fill      = xl_fill("EEF2F8" if is_step else "FFFFFF")
        a.alignment = xl_align("left", "center")
        b.alignment = xl_align("left", "center", wrap=True)
        if label:
            a.border = xl_border()
            b.border = xl_border()

    # ── SHEET 3: CLIENTES REFERENCIA ─────────────────────────
    ws3 = wb.create_sheet("Clientes Referencia")
    ws3.sheet_view.showGridLines = False

    ws3.merge_cells("A1:F1")
    ws3["A1"] = "CLIENTES PRIORITARIOS — ZONA INDUSTRIAL I"
    ws3["A1"].font      = xl_font(bold=True, size=13, color="FFFFFF")
    ws3["A1"].fill      = xl_fill("1A2E4A")
    ws3["A1"].alignment = xl_align("center")
    ws3.row_dimensions[1].height = 22

    col_widths3 = [5, 22, 22, 40, 18, 18]
    for i, w in enumerate(col_widths3, 1):
        ws3.column_dimensions[get_column_letter(i)].width = w

    hdrs3 = ["#", "Empresa", "Tipo", "Por que son importantes", "Telefono", "Email"]
    for ci, h in enumerate(hdrs3, 1):
        c = ws3.cell(row=2, column=ci, value=h)
        c.font      = xl_font(bold=True, size=9, color="FFFFFF")
        c.fill      = xl_fill("C89B2F")
        c.alignment = xl_align("center")
        c.border    = xl_border()
    ws3.row_dimensions[2].height = 18

    clientes3 = [
        ["1", "SASGO", "Fabricante electrico",
         "Fabrica torres y postes para CORPOELEC. ALTA necesidad de galvanizado. Misma zona.",
         "+58 412-536-3346", "importacion@sasgo.com.ve"],
        ["2", "Industrias Marullo, S.A.", "Metal / agroindustrial",
         "Estructuras de acero y maquinaria. Fundada 1955. Muy activa marzo 2026.",
         "+58 424-514-3859", "marullo@hotmail.com"],
        ["3", "GEDISA", "Distribuidor electrico",
         "60 anos en el mercado. Reabrieron enero 2026. En misma zona industrial.",
         "+58 414-503-2546", "gedisa@gedisa.com.ve"],
        ["4", "CORPOELEC Lara — Alm. N01", "Empresa estatal electrica",
         "Compraron 20.000 varillas de tierra a PYGLARA (2004-2005). Almacen en misma Calle 16.",
         "0251-239-4050", "corpoelecresponde@corpoelec.gob.ve"],
        ["5", "MANPEG", "Fabricante postes",
         "Postes hexagonales y electricos hasta 40m. Alta necesidad de galvanizado.",
         "+58 424-561-2759", "@manpeg.ca (IG)"],
        ["6", "GEMACA", "Industrial",
         "Tuberia PEAD. Necesidad de galvanizado baja pero vale verificar.",
         "+58 251-237-6922", "gemaca.com"],
    ]

    for ri, row_data in enumerate(clientes3, 3):
        ws3.row_dimensions[ri].height = 22
        bg = "F4F8FF" if ri % 2 != 0 else "FFFFFF"
        for ci, val in enumerate(row_data, 1):
            c = ws3.cell(row=ri, column=ci, value=val)
            c.font      = xl_font(size=9, bold=(ci==1), color="1A2E4A" if ci==1 else "000000")
            c.fill      = xl_fill(bg)
            c.alignment = xl_align("left", "center", wrap=True)
            c.border    = xl_border()

    out_path = os.path.join(OUTPUT_DIR, "PYGLARA_Registro_Visitas.xlsx")
    wb.save(out_path)
    print(f"EXCEL guardado: {out_path}")


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_word()
    build_excel()
    print("\nListo. Ambos archivos generados en la carpeta drafts/")
