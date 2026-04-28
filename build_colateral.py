"""
Build script for PYGLARA Inventario de Colateral y Valoración de Activos
Produces: 06-bank-package/bank-docs/PYGLARA-Inventario-Colateral-Valoracion-Activos.docx
          06-bank-package/bank-docs/PYGLARA-Inventario-Colateral-Valoracion-Activos.pdf

Pipeline:
  1. pandoc markdown -> docx (reference.docx styles)
  2. python-docx: professional formatting (colors, page breaks, headers/footers, table styling)
  3. Append: proforma invoice image page
  4. Append: labeled asset photo gallery (2-column)
  5. LibreOffice headless: docx -> pdf via c:\tmp\
"""

import os, shutil, subprocess, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

# ─── Paths ────────────────────────────────────────────────────────────────────
ROOT = Path(r"c:\Users\amont\Desktop\PYG")
MD   = ROOT / "_bmad-output/implementation-artifacts/epic-6-bank/inventario-colateral-valoracion-activos.md"
OUT_DOCX = ROOT / "06-bank-package/bank-docs/PYGLARA-Inventario-Colateral-Valoracion-Activos.docx"
OUT_PDF  = ROOT / "06-bank-package/bank-docs/PYGLARA-Inventario-Colateral-Valoracion-Activos.pdf"
TMP_DOCX = Path(r"c:\tmp\pyglara_colateral.docx")
TMP_PDF  = Path(r"c:\tmp\pyglara_colateral.pdf")
INVOICE_PDF = ROOT / "04-investor-document/working-drafts/PROFORMA INVOICE PRENSADOS Y GALVAIZADOS DE LARA.pdf"
INVOICE_IMG = Path(r"c:\tmp\invoice_page1.png")
ASSETS_BROCHURE = ROOT / "assets/brochure-images"
ASSETS_PLANT    = ROOT / "assets/plant-photos"

PANDOC = r"C:\Users\amont\AppData\Local\Pandoc\pandoc.exe"
LIBRE  = r"C:\Program Files\LibreOffice\program\soffice.exe"

# ─── Brand colors ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1A, 0x2E, 0x4A)
GOLD  = RGBColor(0xC8, 0xA0, 0x32)
LGREY = RGBColor(0xF5, 0xF5, 0xF5)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)

def rgb_hex(c: RGBColor) -> str:
    return '%02X%02X%02X' % (c[0], c[1], c[2])

# ─── Section headers that get a page break BEFORE them ───────────────────────
SECTION_STARTS = [
    "SECCIÓN 1",
    "SECCIÓN 2",
    "SECCIÓN 3",
    "SECCIÓN 4",
    "SECCIÓN 5",
    "SECCIÓN 6",
    "DECLARACIÓN",
    "APÉNDICE",
]

# ─── Helper: OxmlElement shortcuts ───────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_borders(cell, color='1A2E4A', sz='6'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def cant_split_row(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit')
    cs.set(qn('w:val'), '1')
    trPr.append(cs)

def keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepNext')
    kwn.set(qn('w:val'), '1')
    pPr.append(kwn)

def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)

def add_field(para, field_code):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into paragraph."""
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)
    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code
    run2._r.append(instrText)
    run3 = para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar_end)

def add_gold_left_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), rgb_hex(GOLD))
    pBdr.append(left)
    pPr.append(pBdr)

def insert_page_break_paragraph(doc, before_para):
    """Insert a paragraph with page break XML before given paragraph."""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_br = OxmlElement('w:br')
    new_br.set(qn('w:type'), 'page')
    new_r.append(new_br)
    new_p.append(new_r)
    before_para._p.addprevious(new_p)

# ─── Step 1: Pandoc ───────────────────────────────────────────────────────────
def run_pandoc():
    print("Running pandoc...")
    os.makedirs(TMP_DOCX.parent, exist_ok=True)
    result = subprocess.run([
        PANDOC, str(MD),
        "-o", str(TMP_DOCX),
        "--from", "markdown",
        "--to", "docx",
    ], capture_output=True, text=True)
    if result.returncode != 0:
        print("PANDOC ERROR:", result.stderr)
        sys.exit(1)
    print("  pandoc done -> " + str(TMP_DOCX))

# ─── Step 2: Format document ─────────────────────────────────────────────────
def format_doc():
    print("Formatting document...")
    doc = Document(str(TMP_DOCX))

    # Page margins
    for section in doc.sections:
        section.page_width  = Cm(21.59)
        section.page_height = Cm(27.94)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK

    # Heading styles
    for h_name, sz, bold, color in [
        ('Heading 1', 16, True, NAVY),
        ('Heading 2', 13, True, NAVY),
        ('Heading 3', 11, True, NAVY),
    ]:
        try:
            s = doc.styles[h_name]
            s.font.name = 'Calibri'
            s.font.size = Pt(sz)
            s.font.bold = bold
            s.font.color.rgb = color
        except KeyError:
            pass

    # Build ordered list of body elements
    body = doc.element.body
    para_map  = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    # ── Pass 1: headings, page breaks, borders ────────────────────────────────
    prev_para = None
    for elem in body:
        if elem.tag == qn('w:p'):
            para = para_map.get(elem)
            if para is None:
                continue
            sname = para.style.name if para.style else ''

            if sname.startswith('Heading'):
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after  = Pt(4)
                keep_with_next(para)

                if sname == 'Heading 2':
                    add_gold_left_border(para)
                    text = para.text.strip().upper()
                    if any(text.startswith(s) for s in SECTION_STARTS):
                        page_break_before(para)

            elif sname in ('Normal', 'Body Text', 'First Paragraph', ''):
                para.paragraph_format.space_after = Pt(6)

            prev_para = para

        elif elem.tag == qn('w:tbl'):
            tbl = table_map.get(elem)
            if tbl is None:
                continue
            # Style table
            _style_table(tbl)
            # Keep heading before table together
            if prev_para is not None:
                keep_with_next(prev_para)

    # ── Cover page: make title huge navy, center ──────────────────────────────
    # First Heading 1 is the document title
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 1':
            for run in para.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = NAVY
                run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Cm(2)
            para.paragraph_format.space_after  = Cm(0.5)
            break

    # ── Headers & footers ─────────────────────────────────────────────────────
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("PYGLARA — Inventario de Colateral y Valoración de Activos")
        r.font.name = 'Calibri'
        r.font.size = Pt(8)
        r.font.color.rgb = NAVY
        r.font.italic = True
        # gold rule under header
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), rgb_hex(GOLD))
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fl = fp.add_run("Prensados y Galvanizados de Lara, S.A.  |  RIF: J-07014488-0  |  Barquisimeto, Lara  |  Página ")
        fl.font.name = 'Calibri'
        fl.font.size = Pt(8)
        fl.font.color.rgb = NAVY
        add_field(fp, ' PAGE ')
        fr = fp.add_run(" de ")
        fr.font.name = 'Calibri'
        fr.font.size = Pt(8)
        fr.font.color.rgb = NAVY
        add_field(fp, ' NUMPAGES ')

    doc.save(str(TMP_DOCX))
    print(f"  formatting done -> {TMP_DOCX}")

def _style_table(tbl):
    """Apply navy border and alternating-row styling to a table."""
    hdr_done = False
    for i, row in enumerate(tbl.rows):
        cant_split_row(row)
        is_header = not hdr_done
        for j, cell in enumerate(row.cells):
            if is_header:
                set_cell_bg(cell, NAVY)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
            elif i % 2 == 0:
                set_cell_bg(cell, LGREY)
            else:
                set_cell_bg(cell, WHITE)

            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                if not is_header:
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        run.font.color.rgb = DARK

            set_cell_borders(cell, rgb_hex(NAVY), '4')
        hdr_done = True

# ─── Step 3: Convert invoice PDF to image ─────────────────────────────────────
def convert_invoice_to_image():
    print("Converting proforma invoice to image...")
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(INVOICE_PDF))
        page = doc[0]
        mat = fitz.Matrix(2.0, 2.0)  # 2x zoom = ~144 dpi
        pix = page.get_pixmap(matrix=mat)
        pix.save(str(INVOICE_IMG))
        doc.close()
        print(f"  invoice image -> {INVOICE_IMG}")
        return True
    except Exception as e:
        print(f"  WARNING: Could not convert invoice: {e}")
        return False

# ─── Step 4: Append invoice page ─────────────────────────────────────────────
def append_invoice(doc):
    print("  Appending invoice page...")
    # Page break
    pb_para = doc.add_paragraph()
    pb_para.paragraph_format.space_before = Pt(0)
    pb_para.paragraph_format.space_after  = Pt(0)
    run = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

    # Section title
    h = doc.add_heading('APÉNDICE II — COTIZACIÓN DE ZINC SHG 99,995%', level=2)
    h.style.font.color.rgb = NAVY
    add_gold_left_border(h)

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("Pan American Zinc LLC — Proforma Invoice N° 02/20/2026")
    sr.font.name = 'Calibri'
    sr.font.size = Pt(10)
    sr.font.bold = True
    sr.font.color.rgb = NAVY
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    detail = doc.add_paragraph()
    dr = detail.add_run("Zinc SHG 99,995% — USD $3.930/TM — 25 TM — Total: USD $98.250")
    dr.font.name = 'Calibri'
    dr.font.size = Pt(9)
    dr.font.italic = True
    dr.font.color.rgb = DARK
    detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detail.paragraph_format.space_after = Pt(10)

    # Invoice image
    if INVOICE_IMG.exists():
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(INVOICE_IMG), width=Cm(15))
    else:
        doc.add_paragraph("[Proforma invoice image not available — see attached PDF]")

# ─── Step 5: Append photo gallery ─────────────────────────────────────────────
# Curated asset photos: (label, caption, path)
ASSET_PHOTOS = [
    (
        "Galpones Industriales — Zona Industrial I, Barquisimeto",
        "Vista frontal de los dos galpones. Superficie total: ~3.640 m².",
        ASSETS_BROCHURE / "facade.jpg",
    ),
    (
        "Interior Galpón Principal — Puentes Grúa",
        "Nave principal (2.500 m²). Puentes grúa de 2 TM y 5 TM visibles.",
        ASSETS_BROCHURE / "hall.jpg",
    ),
    (
        "Cuba de Galvanizado 7 Metros — W. Pilling Riepe (Alemania)",
        "Cuba principal instalada. Quemadores en buen estado. Requiere carga de zinc (50 TM).",
        ASSETS_BROCHURE / "kettle.jpg",
    ),
    (
        "Línea de Cobre Electrolítico",
        "Tanques de proceso y rectificadores trifásicos. Capacidad: 936 varillas/día.",
        ASSETS_BROCHURE / "copper.jpg",
    ),
    (
        "Producto Galvanizado — Pletinas (Barras Planas)",
        "Muestra de producción. Récord 2015: 675 TM de pletinas galvanizadas.",
        ASSETS_BROCHURE / "product.jpg",
    ),
    (
        "Producto Galvanizado — Soportes y Herrajes",
        "Piezas galvanizadas listas para entrega a cliente.",
        ASSETS_BROCHURE / "bracket.jpg",
    ),
    (
        "Horno de Galvanizado — Placa W. Pilling Riepe",
        "Identificación del fabricante alemán en el horno de galvanizado.",
        ASSETS_PLANT / "WhatsApp Image 2026-04-01 at 1.27.37 PM.jpeg",
    ),
    (
        "Cuba de Galvanizado 7m — Vista Operacional",
        "Vista del área de operación de la cuba principal con sistema de grúas.",
        ASSETS_PLANT / "WhatsApp Image 2026-04-01 at 1.27.38 PM.jpeg",
    ),
    (
        "Producto en Patio — Acero Galvanizado",
        "Producto terminado en patio de la planta.",
        ASSETS_PLANT / "WhatsApp Image 2026-04-01 at 1.25.24 PM.jpeg",
    ),
    (
        "Operaciones de Despacho",
        "Carga de acero galvanizado para entrega a cliente.",
        ASSETS_PLANT / "WhatsApp Image 2026-04-01 at 1.25.26 PM.jpeg",
    ),
]

def append_photo_gallery(doc):
    print("  Appending photo gallery...")
    pb_para = doc.add_paragraph()
    run = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

    h = doc.add_heading('APÉNDICE III — REGISTRO FOTOGRÁFICO DE ACTIVOS', level=2)
    h.style.font.color.rgb = NAVY
    add_gold_left_border(h)

    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Las siguientes fotografías muestran los activos incluidos en el inventario de colateral. "
        "Fotografías adicionales disponibles a solicitud de la institución financiera."
    )
    ir.font.name = 'Calibri'
    ir.font.size = Pt(9)
    ir.font.italic = True
    ir.font.color.rgb = DARK
    intro.paragraph_format.space_after = Pt(10)

    # 2-column gallery table
    # Filter to existing images only
    photos = [(lbl, cap, path) for lbl, cap, path in ASSET_PHOTOS if path.exists()]
    if not photos:
        doc.add_paragraph("[No photo files found — add images to assets/ folder]")
        return

    # Build pairs
    pairs = []
    for i in range(0, len(photos), 2):
        left = photos[i]
        right = photos[i+1] if i+1 < len(photos) else None
        pairs.append((left, right))

    for left, right in pairs:
        tbl = doc.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_names = [s.name for s in doc.styles]
        if 'Table Grid' in style_names:
            tbl.style = doc.styles['Table Grid']

        # Set column widths
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), '4320')  # ~7.6cm in twentieths of a point
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

        # Image row
        img_row = tbl.rows[0]
        cant_split_row(img_row)
        for j, item in enumerate([left, right]):
            cell = img_row.cells[j]
            set_cell_bg(cell, WHITE)
            _remove_cell_borders(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if item:
                try:
                    run = para.add_run()
                    run.add_picture(str(item[2]), width=Cm(7.8))
                except Exception as e:
                    para.add_run(f"[Image error: {e}]")

        # Caption row
        cap_row = tbl.rows[1]
        cant_split_row(cap_row)
        for j, item in enumerate([left, right]):
            cell = cap_row.cells[j]
            set_cell_bg(cell, LGREY)
            _remove_cell_borders(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(6)
            if item:
                r1 = para.add_run(item[0] + "\n")
                r1.font.name = 'Calibri'
                r1.font.size = Pt(8)
                r1.font.bold = True
                r1.font.color.rgb = NAVY
                r2 = para.add_run(item[1])
                r2.font.name = 'Calibri'
                r2.font.size = Pt(7.5)
                r2.font.color.rgb = DARK

        # Spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)

def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

# ─── Step 6: Export via LibreOffice ──────────────────────────────────────────
def export_pdf():
    print("Exporting PDF via LibreOffice...")
    # Kill any lingering soffice
    subprocess.run(['taskkill', '/F', '/IM', 'soffice.exe'], capture_output=True)
    result = subprocess.run([
        LIBRE, '--headless', '--convert-to', 'pdf',
        '--outdir', str(TMP_DOCX.parent),
        str(TMP_DOCX),
    ], capture_output=True, text=True, timeout=120)
    print("  LO stdout:", result.stdout.strip())
    if result.returncode != 0:
        print("  LO stderr:", result.stderr.strip())
    if TMP_PDF.exists():
        shutil.copy(str(TMP_PDF), str(OUT_PDF))
        print(f"  PDF -> {OUT_PDF}")
    else:
        print("  ERROR: PDF not created. Check LibreOffice output above.")

# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)

    # 1. pandoc
    run_pandoc()

    # 2. format
    format_doc()

    # 3. invoice image
    have_invoice = convert_invoice_to_image()

    # 4+5. append appendices
    doc = Document(str(TMP_DOCX))
    append_invoice(doc)
    append_photo_gallery(doc)
    doc.save(str(TMP_DOCX))
    print(f"  appendices added -> {TMP_DOCX}")

    # Copy docx to output
    shutil.copy(str(TMP_DOCX), str(OUT_DOCX))
    print(f"  DOCX -> {OUT_DOCX}")

    # 6. pdf
    export_pdf()

    print("\nDone.")
    print(f"  DOCX: {OUT_DOCX}")
    print(f"  PDF:  {OUT_PDF}")

if __name__ == '__main__':
    main()
